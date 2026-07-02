#!/usr/bin/env python3
"""Build expanded Word table: portal sessões + atas + YouTube cross-reference."""

from __future__ import annotations

import json
import re
import subprocess
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

UA = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
OUT = Path("/workspace/TCU_Sessoes_Julgamento_Expandido_2025_2026.docx")
CACHE_SESSOES = Path("/tmp/sessoes.html")
CACHE_PORTAL_JSON = Path("/tmp/portal_sessions.json")

COLEGIADO_BY_HOUR = {
    "10:00": "Plenário",
    "10:30": "2ª Câmara",
    "11:00": "Plenário",
    "14:30": "Plenário",
    "15:00": "1ª Câmara",
}
COLEGIADO_BY_COL = {"p": "Plenário", "1c": "1ª Câmara", "2c": "2ª Câmara"}


def fetch(url: str, timeout: int = 30) -> str:
    req = urllib.request.Request(url, headers=UA)
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        return resp.read().decode("utf-8", errors="replace")


def parse_rsc(html: str) -> str:
    chunks = re.findall(r'self\.__next_f\.push\(\[1,"(.+?)"\]\)', html, re.DOTALL)
    return "".join(chunks).encode("utf-8").decode("unicode_escape", errors="replace")


def fix_mojibake(text: str) -> str:
    if "Ã" in text or "â" in text:
        try:
            return text.encode("latin1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
    return text


def load_portal_listing() -> list[dict]:
    """Parse cached portal /sessoes HTML (live fetch often blocked for bots)."""
    html_path = CACHE_SESSOES
    if html_path.exists():
        html = html_path.read_text(encoding="utf-8", errors="replace")
    else:
        html = fetch("https://portal.tcu.gov.br/sessoes")
        if "Acesso Bloqueado" in html:
            raise RuntimeError("Portal bloqueado e cache ausente.")
        html_path.write_text(html, encoding="utf-8")

    full = parse_rsc(html)
    ata_by_num: dict[str, dict] = {}
    for n, ano, col in re.findall(r"#/doc/ata-sessao/(\d+)/(\d{4})/([^/]+)/publica", full):
        ata_by_num[f"{n}/{ano}/{col}"] = {
            "ata_url": f"https://pesquisa.apps.tcu.gov.br/#/doc/ata-sessao/{n}/{ano}/{col}/publica",
            "numeroAta": int(n),
            "ano": ano,
            "col": col,
            "colegiado": COLEGIADO_BY_COL.get(col, col),
        }

    # Map session id -> date/time from nearby HTML context
    rows: dict[str, dict] = {}
    for m in re.finditer(r"/sessoes/sessao/(\d+)", full):
        sid = m.group(1)
        ctx = full[max(0, m.start() - 600) : m.start() + 120]
        dm = re.search(r"(\d{2}/\d{2}/202[56])", ctx)
        hm = re.search(r"(\d{2}:\d{2})", ctx)
        if not dm:
            continue
        date_br = dm.group(1)
        hora = hm.group(1) if hm else ""
        d, mo, y = date_br.split("/")
        rows[sid] = {
            "portal_id": sid,
            "portal_url": f"https://portal.tcu.gov.br/sessoes/sessao/{sid}",
            "data_br": date_br,
            "data_iso": f"{y}-{mo}-{d}",
            "hora": hora,
            "ano": y,
            "colegiado": COLEGIADO_BY_HOUR.get(hora, ""),
            "naturezaSessao": "Ordinária",
        }

    # Enrich with ata data from earlier fetch json if available
    if CACHE_PORTAL_JSON.exists():
        for item in json.loads(CACHE_PORTAL_JSON.read_text(encoding="utf-8")):
            sid = str(item.get("session_id", ""))
            if sid not in rows:
                continue
            if item.get("ata_url"):
                rows[sid]["linkBaixarAta"] = item["ata_url"]
                m = re.search(r"/(\d+)/(202[56])/([^/]+)/publica", item["ata_url"])
                if m:
                    n, ano, col = m.groups()
                    rows[sid]["numeroAta"] = int(n)
                    rows[sid]["colegiado"] = COLEGIADO_BY_COL.get(col, rows[sid].get("colegiado", ""))

    # Special session types by date/time heuristics
    for row in rows.values():
        if row["data_br"] == "10/06/2026" and row["hora"] == "10:00":
            row["naturezaSessao"] = "Contas do Presidente da República"
            row["colegiado"] = "Plenário"
        elif row["data_br"] == "20/05/2026" and row["hora"] == "14:30":
            row["naturezaSessao"] = "Posse de Ministro"
            row["colegiado"] = "Plenário"
        elif row["data_br"] == "19/05/2026" and row["hora"] == "14:30":
            row["naturezaSessao"] = "Extraordinária"
            row["colegiado"] = "Plenário"

    return sorted(rows.values(), key=lambda x: (x["data_iso"], x["hora"]))


def load_youtube_sessions() -> list[dict]:
    items: list[dict] = []
    seen: set[str] = set()
    for tab in ("streams", "videos"):
        r = subprocess.run(
            ["yt-dlp", "--flat-playlist", "-j", f"https://www.youtube.com/@TCUoficial/{tab}"],
            capture_output=True,
            text=True,
            check=False,
        )
        for line in r.stdout.splitlines():
            if not line.strip():
                continue
            v = json.loads(line)
            if v["id"] in seen:
                continue
            seen.add(v["id"])
            items.append(
                {"title": v["title"], "id": v["id"], "url": f"https://www.youtube.com/watch?v={v['id']}"}
            )
    return items


def parse_yt_date(title: str) -> str | None:
    m = re.search(r"(\d{2})/(\d{2})/(\d{4})", title)
    if not m:
        return None
    a, b, c = map(int, m.groups())
    if a > 12:
        return f"{c:04d}-{b:02d}-{a:02d}"
    if b > 12:
        return f"{c:04d}-{a:02d}-{b:02d}"
    return f"{c:04d}-{b:02d}-{a:02d}"


def colegiado_from_title(title: str) -> str | None:
    t = title.lower()
    if re.search(r"2nd|2ª|segunda|2a c", t):
        return "2ª Câmara"
    if re.search(r"1st|1ª|primeira|1a c", t):
        return "1ª Câmara"
    if "plen" in t or "plenary" in t:
        return "Plenário"
    return None


def match_youtube(session: dict, yt_items: list[dict]) -> dict | None:
    col = session.get("colegiado", "")
    date_iso = session["data_iso"]
    for yt in yt_items:
        yd = parse_yt_date(yt["title"])
        if yd != date_iso:
            continue
        ycol = colegiado_from_title(yt["title"])
        if ycol and col and (ycol[0] == col[0] or ycol in col or col in ycol):
            return yt
    for yt in yt_items:
        if parse_yt_date(yt["title"]) == date_iso and colegiado_from_title(yt["title"]) == col:
            return yt
    for yt in yt_items:
        if parse_yt_date(yt["title"]) == date_iso:
            return yt
    return None


def classify_row(session: dict, yt_match: dict | None) -> dict:
    video_link = ""
    tipo = "—"
    status = "Somente ata / vídeo não localizado"
    fonte = "Portal TCU (cache HTML)"

    if yt_match:
        video_link = yt_match["url"]
        tipo = "YouTube"
        status = "Público (canal @TCUoficial)"
        fonte = "Portal TCU + YouTube"

    titulo = f"{session.get('colegiado', '')} {session.get('naturezaSessao', '')} — {session.get('data_br', '')}"
    if session.get("hora"):
        titulo += f" {session['hora']}"
    if yt_match:
        titulo = yt_match["title"]

    return {
        "data": session.get("data_br", ""),
        "hora": session.get("hora", ""),
        "titulo": titulo.strip(),
        "video_link": video_link,
        "ata_link": session.get("linkBaixarAta", ""),
        "portal_link": session.get("portal_url", "https://portal.tcu.gov.br/sessoes"),
        "fonte": fonte,
        "tipo": tipo,
        "status": status,
        "ano": session.get("ano", ""),
        "colegiado": session.get("colegiado", ""),
    }


def add_hyperlink(paragraph, url: str, text: str | None = None) -> None:
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main() -> None:
    print("Carregando listagem portal (cache)...")
    portal_rows = load_portal_listing()
    print(f"  {len(portal_rows)} sessões no cache portal (2025-2026)")

    print("Carregando YouTube...")
    yt_items = load_youtube_sessions()

    rows: list[dict] = []
    matched_yt: set[str] = set()

    for s in portal_rows:
        if s["ano"] not in ("2025", "2026"):
            continue
        yt = match_youtube(s, yt_items)
        if yt:
            matched_yt.add(yt["id"])
        rows.append(classify_row(s, yt))

    # YouTube-only (incl. 2025) not matched to portal cache
    for yt in yt_items:
        yd = parse_yt_date(yt["title"])
        if not yd or yd[:4] not in ("2025", "2026"):
            continue
        if not re.search(
            r"plen|câmara|camara|chamber|contas do|ordin|extraordin|posse|inauguration",
            yt["title"],
            re.I,
        ):
            continue
        if yt["id"] in matched_yt:
            continue
        rows.append(
            {
                "data": datetime.strptime(yd, "%Y-%m-%d").strftime("%d/%m/%Y"),
                "hora": "",
                "titulo": yt["title"],
                "video_link": yt["url"],
                "ata_link": "",
                "portal_link": "https://portal.tcu.gov.br/sessoes",
                "fonte": "YouTube (sem match no cache portal)",
                "tipo": "YouTube",
                "status": "Público (canal @TCUoficial)",
                "ano": yd[:4],
                "colegiado": colegiado_from_title(yt["title"]) or "",
            }
        )

    rows.sort(key=lambda r: (r["ano"], r["data"], r["hora"], r["titulo"]))

    doc = Document()
    title = doc.add_heading("Sessões de Julgamento TCU 2025–2026 — Tabela Expandida", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Canal: https://www.youtube.com/@TCUoficial")
    doc.add_paragraph("Portal: https://portal.tcu.gov.br/sessoes")
    doc.add_paragraph("Atas: https://pesquisa.apps.tcu.gov.br/#/pesquisa/ata-sessao")
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph(f"Total de registros: {len(rows)}")

    note = doc.add_paragraph()
    note.add_run("Metodologia: ").bold = True
    note.add_run(
        "Cruzamento entre cache HTML do portal /sessoes (25 sessões recentes com data, colegiado inferido "
        "por horário/ata), atas em pesquisa.apps.tcu.gov.br e replays públicos no YouTube. "
        "Coluna Vídeo do portal (incl. unlisted e Teams) exige navegador — não disponível por bloqueio WAF "
        "a bots neste ambiente. Teams: transmissão paralela desde 2022; links em /sessoes quando publicados."
    )

    headers = ["#", "Data", "Hora", "Colegiado", "Título", "Link vídeo", "Link ata", "Portal", "Fonte", "Tipo", "Status"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    for idx, row in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = row["data"]
        cells[2].text = row.get("hora", "") or "—"
        cells[3].text = row["colegiado"] or "—"
        cells[4].text = row["titulo"][:220]
        cells[5].text = row["video_link"] or "—"
        cells[6].text = "Ata" if row["ata_link"] else "—"
        cells[7].text = row["portal_link"]
        cells[8].text = row["fonte"]
        cells[9].text = row["tipo"]
        cells[10].text = row["status"]
        if row["video_link"]:
            p = cells[5].paragraphs[0]
            p.clear()
            add_hyperlink(p, row["video_link"], "YouTube")
        if row["ata_link"]:
            p = cells[6].paragraphs[0]
            p.clear()
            add_hyperlink(p, row["ata_link"], "Ata")

    # Summary sheet as paragraph
    from collections import Counter

    doc.add_paragraph()
    p = doc.add_paragraph("Resumo: ")
    for ano in sorted({r["ano"] for r in rows}):
        sub = [r for r in rows if r["ano"] == ano]
        com_video = sum(1 for r in sub if r["video_link"])
        com_ata = sum(1 for r in sub if r["ata_link"])
        p.add_run(f"{ano}: {len(sub)} registros ({com_video} com vídeo, {com_ata} com ata); ")

    doc.save(OUT)
    print(f"Salvo: {OUT} ({len(rows)} linhas)")
    json.dump(rows, open("/tmp/tcu_expandido.json", "w"), ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
